"""
CV DOCX Generator (daily-cv skill).

Generates a tuned, single-page CV as a .docx via python-docx.
Per R19: use python-docx (not a JS heredoc) to avoid multi-byte encoding errors.

ALL person-specific content is loaded from the profile (profile/PROFILE.json,
falling back to profile/PROFILE.example.json). Nothing about any person is
hardcoded here. See profile_loader.py and profile/README.md.

USAGE:
  python generate_cv.py --company Acme --role "VP Engineering" --run-date 20260101
  # optional: --profile <path>  --summary "..."  --output-dir ./out
The summary/competencies come from the profile but can be overridden per run with
--summary / --competencies to match a job description's keywords.
"""
import argparse
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from profile_loader import (
    load_profile,
    add_target_args,
    resolve_target,
    contact_line,
    competencies_text,
)

DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)


def parse_args():
    p = argparse.ArgumentParser(description="Generate a single-page CV DOCX from the profile.")
    add_target_args(p)
    p.add_argument("--summary", help="Override the profile summary for this run (JD-tuned).")
    p.add_argument("--competencies", help="Override competencies (pipe- or comma-separated).")
    return p.parse_args()


def clear_para(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def simple_para(doc, text, bold=False, italic=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
                sp_before=0, sp_after=0, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
    return p


def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def role_header(doc, title, company, dates, city=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.name = 'Calibri'
    r2 = p.add_run('  |  ')
    r2.font.size = Pt(9)
    r2.font.name = 'Calibri'
    r3 = p.add_run(company)
    r3.bold = True
    r3.italic = True
    r3.font.size = Pt(9)
    r3.font.name = 'Calibri'
    r4 = p.add_run('  |  ' + dates)
    r4.font.size = Pt(8.5)
    r4.font.name = 'Calibri'
    r4.font.color.rgb = DARK_GRAY
    if city:
        r5 = p.add_run('  |  ' + city)
        r5.font.size = Pt(8.5)
        r5.font.name = 'Calibri'
        r5.font.color.rgb = DARK_GRAY
    return p


def bullet(doc, text, size=8.5, sp_after=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    rb = p.add_run(u'•  ')
    rb.font.size = Pt(size)
    rb.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    return p


def build(profile, target, summary, competencies):
    doc = Document()

    normal = doc.styles['Normal']
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.font.size = Pt(9)
    normal.font.name = 'Calibri'

    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.3)
    sec.bottom_margin = Inches(0.3)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    ident = profile.get("identity", {})

    # NAME BLOCK (R9 — exact display_name from the profile)
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(ident["display_name"])
    r_name.bold = True
    r_name.font.size = Pt(26)
    r_name.font.name = 'Calibri'
    r_name.font.color.rgb = DARK_BLUE

    simple_para(doc, contact_line(profile), size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
                sp_after=1, color=DARK_GRAY)

    section_heading(doc, 'Professional Summary')
    simple_para(doc, summary, size=8.5, sp_after=0)

    section_heading(doc, 'Core Competencies')
    simple_para(doc, competencies, size=8.5, sp_after=0)

    section_heading(doc, 'Professional Experience')
    for role in profile.get("experience", []):
        role_header(doc, role.get("title", ""), role.get("company", ""),
                    role.get("dates", ""), role.get("city", ""))
        bullets = role.get("bullets", [])
        for i, b in enumerate(bullets):
            bullet(doc, b, sp_after=1 if i == len(bullets) - 1 else 1.5)

    earlier = profile.get("earlier_career", [])
    if earlier:
        section_heading(doc, 'Earlier Career')
        simple_para(doc, '  |  '.join(earlier), size=8.5, sp_after=0)

    education = profile.get("education", [])
    if education:
        section_heading(doc, 'Education')
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(0)
        p_edu.paragraph_format.space_after = Pt(0)
        for idx, edu in enumerate(education):
            r_e = p_edu.add_run(edu.get("degree", ""))
            r_e.bold = True
            r_e.font.size = Pt(8.5)
            r_e.font.name = 'Calibri'
            tail = f" | {edu.get('institution', '')} | {edu.get('dates', '')}"
            if edu.get("note"):
                tail += f" ({edu['note']})"
            if idx < len(education) - 1:
                tail += "      "
            p_edu.add_run(tail).font.size = Pt(8.5)

    volunteering = profile.get("volunteering", [])
    if volunteering:
        section_heading(doc, 'Volunteering')
        simple_para(doc, '  |  '.join(volunteering), size=8.5, sp_after=0)

    out_path = os.path.join(target["output_dir"],
                            f"{ident.get('filename_slug', 'CV')}_CV_{target['company']}_{target['run_date']}.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  (target: {target['role']} @ {target['company']})")
    print(f"Size: {os.path.getsize(out_path):,} bytes")
    print(f"Profile: {profile.get('_source_path')}")


def main():
    args = parse_args()
    profile = load_profile(args.profile)
    target = resolve_target(profile, args)
    summary = args.summary or profile.get("summary", "")
    if args.competencies:
        competencies = args.competencies.replace(",", "  |  ")
    else:
        competencies = competencies_text(profile)
    build(profile, target, summary, competencies)


if __name__ == "__main__":
    main()
