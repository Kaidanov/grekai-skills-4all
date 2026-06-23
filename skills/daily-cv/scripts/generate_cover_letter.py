"""
Cover Letter DOCX Generator (daily-cv skill).

Generates a single-page, tailored cover letter as a .docx via python-docx.

The header (name + contact) is loaded from the profile (profile/PROFILE.json,
falling back to profile/PROFILE.example.json). Nothing about any person is
hardcoded here. See profile_loader.py and profile/README.md.

The BODY of the letter is run-specific — the model writes 2-3 paragraphs from the
JD + company intel. Pass them via:
  --body-file <path>   (one paragraph per blank-line-separated block), or
  --body "para1" --body "para2" ...
If omitted, a neutral placeholder body is generated so the script still runs.

R18 salutation: use "Dear <FirstName>," when a recruiter name is known, otherwise
"Dear Hiring Team, <Company>,". Override with --salutation.
"""
import argparse
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from profile_loader import load_profile, add_target_args, resolve_target, contact_line

DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)


def parse_args():
    p = argparse.ArgumentParser(description="Generate a single-page cover letter DOCX from the profile.")
    add_target_args(p)
    p.add_argument("--body", action="append", help="A body paragraph (repeatable).")
    p.add_argument("--body-file", dest="body_file",
                   help="Path to a text file; paragraphs separated by blank lines.")
    return p.parse_args()


def read_body(args, target, profile):
    if args.body:
        return args.body
    if args.body_file:
        with open(args.body_file, "r", encoding="utf-8") as fh:
            blocks = [b.strip().replace("\n", " ") for b in fh.read().split("\n\n")]
        return [b for b in blocks if b]
    # Placeholder body so the script runs end-to-end. The model replaces these.
    headline = profile.get("identity", {}).get("headline", "")
    return [
        (f"This letter expresses interest in the {target['role']} role at {target['company']}. "
         "Replace this placeholder paragraph with a specific, source-backed hook about the company "
         "and why the timing is relevant."),
        (f"Background as {headline} maps to this role. Replace this paragraph with 2-3 of the most "
         "relevant achievements from the profile, mapped to the job description's requirements - "
         "using only verified facts and numbers from key_metrics."),
        ("Available for an initial conversation at your convenience. Replace this paragraph with a "
         "concrete call to action."),
    ]


def para(doc, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
         sp_before=0, sp_after=0, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after = Pt(sp_after)
    p.paragraph_format.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
    return p


def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '1F497D')
    pBdr.append(top)
    pPr.append(pBdr)


def build(profile, target, body_paragraphs):
    doc = Document()

    normal = doc.styles['Normal']
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.font.size = Pt(10)
    normal.font.name = 'Calibri'

    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    ident = profile.get("identity", {})

    # NAME (R9)
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_name.add_run(ident["display_name"])
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_BLUE

    # Contact line without the clearance note (kept off cover letters).
    para(doc, contact_line(profile, include_clearance=False), size=8,
         align=WD_ALIGN_PARAGRAPH.CENTER, sp_after=0, color=DARK_GRAY)
    add_hrule(doc)

    para(doc, target["letter_date"], size=9, sp_after=12, color=DARK_GRAY)
    para(doc, target["salutation"], bold=True, size=10, sp_after=8)

    for text in body_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    para(doc, 'Warm regards,', size=10, sp_before=6, sp_after=4)
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after = Pt(0)
    r_sig = p_sig.add_run(ident["display_name"])
    r_sig.bold = True
    r_sig.font.size = Pt(10)
    r_sig.font.name = 'Calibri'
    r_sig.font.color.rgb = DARK_BLUE

    out_path = os.path.join(
        target["output_dir"],
        f"{ident.get('filename_slug', 'CoverLetter')}_CoverLetter_{target['company']}_{target['run_date']}.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  (target: {target['role']} @ {target['company']})")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")
    print(f"Profile: {profile.get('_source_path')}")


def main():
    args = parse_args()
    profile = load_profile(args.profile)
    target = resolve_target(profile, args)
    body = read_body(args, target, profile)
    build(profile, target, body)


if __name__ == "__main__":
    main()
